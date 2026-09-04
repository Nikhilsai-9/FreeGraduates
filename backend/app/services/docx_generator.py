"""DOCX export — structured resume → ATS-friendly Word document.

Layer C of the source repository's rule system forbids tables, columns,
images, and special formatting in the final ATS-safe resume. We obey
that here.

Layer A forbids emojis and special characters in user-visible text.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.engine.schemas import ResumeStructured


def render_resume_to_docx(resume: ResumeStructured) -> bytes:
    """Build a .docx file from a structured resume and return it as bytes."""
    doc = Document()

    # Tight margins for ATS readability (Layer C reference).
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    _set_base_style(doc)

    _render_header(doc, resume)
    _render_summary(doc, resume)
    _render_skills(doc, resume)
    _render_experience(doc, resume)
    _render_education(doc, resume)
    if resume.languages:
        _render_languages(doc, resume)
    if resume.optional_sections:
        for sec in resume.optional_sections:
            _render_optional(doc, sec.title, sec.items)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- Section renderers ----------


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _render_header(doc: Document, resume: ResumeStructured) -> None:
    h = resume.header

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(h.full_name or "")
    name_run.bold = True
    name_run.font.size = Pt(18)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(h.title or "")
    title_run.font.size = Pt(12)
    title_run.italic = True

    if h.contacts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(h.contacts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1C, 0x6A)


def _render_summary(doc: Document, resume: ResumeStructured) -> None:
    if not (resume.summary and resume.summary.summary_text):
        return
    _section_heading(doc, "Summary")
    doc.add_paragraph(resume.summary.summary_text)


def _render_skills(doc: Document, resume: ResumeStructured) -> None:
    if not resume.skills.groups:
        return
    _section_heading(doc, "Skills")
    for grp in resume.skills.groups:
        if not grp.items:
            continue
        line = grp.group_name + ": " + ", ".join(grp.items)
        doc.add_paragraph(line)


def _render_experience(doc: Document, resume: ResumeStructured) -> None:
    if not resume.experience:
        return
    _section_heading(doc, "Experience")
    for exp in resume.experience:
        head = doc.add_paragraph()
        role_run = head.add_run(f"{exp.role} — {exp.company}")
        role_run.bold = True
        head.add_run(f"   {exp.start_date} – {exp.end_date}").italic = True
        for h in exp.highlights:
            bullet = doc.add_paragraph(h, style="List Bullet")
            bullet.paragraph_format.left_indent = Pt(18)


def _render_education(doc: Document, resume: ResumeStructured) -> None:
    if not resume.education:
        return
    _section_heading(doc, "Education")
    for ed in resume.education:
        line = f"{ed.institution} — {ed.degree}"
        if ed.field:
            line += f", {ed.field}"
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.bold = True
        if ed.start_date or ed.end_date:
            para.add_run(f"   {ed.start_date or ''} – {ed.end_date or ''}").italic = True
        if ed.grade:
            doc.add_paragraph(f"Grade: {ed.grade}")


def _render_languages(doc: Document, resume: ResumeStructured) -> None:
    if not resume.languages:
        return
    _section_heading(doc, "Languages")
    doc.add_paragraph(
        ", ".join(f"{l.name} ({l.level})" for l in resume.languages)
    )


def _render_optional(doc: Document, title: str, items: list[str]) -> None:
    if not items:
        return
    _section_heading(doc, title)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
